"""Document loader for various file formats."""

import io
from pathlib import Path
from typing import Dict, List, Optional, Union

import PyPDF2
import pypdf
from docx import Document as DocxDocument
from loguru import logger

from ..utils import sanitize_filename, format_file_size


class DocumentLoader:
    """Load and extract text from various document formats."""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.docx', '.doc'}
    
    def __init__(self):
        self.logger = logger.bind(component="DocumentLoader")
    
    def load_document(self, file_path: Union[str, Path, io.BytesIO], filename: Optional[str] = None) -> Dict[str, str]:
        """
        Load a document and extract its text content.
        
        Args:
            file_path: Path to the document file or BytesIO object
            filename: Original filename (required when using BytesIO)
            
        Returns:
            Dictionary with document metadata and content
        """
        if isinstance(file_path, io.BytesIO):
            if not filename:
                raise ValueError("Filename is required when loading from BytesIO")
            return self._load_from_bytes(file_path, filename)
        
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        self.logger.info(f"Loading document: {file_path.name}")
        
        extension = file_path.suffix.lower()
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format: {extension}")
        
        try:
            if extension == '.pdf':
                content = self._load_pdf(file_path)
            elif extension == '.txt':
                content = self._load_text(file_path)
            elif extension in ['.docx', '.doc']:
                content = self._load_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {extension}")
            
            file_size = file_path.stat().st_size
            
            return {
                'filename': file_path.name,
                'content': content,
                'file_size': file_size,
                'file_size_formatted': format_file_size(file_size),
                'extension': extension,
                'char_count': len(content),
                'word_count': len(content.split()),
            }
            
        except Exception as e:
            self.logger.error(f"Error loading document {file_path.name}: {str(e)}")
            raise
    
    def _load_from_bytes(self, file_bytes: io.BytesIO, filename: str) -> Dict[str, str]:
        """Load document from BytesIO object."""
        extension = Path(filename).suffix.lower()
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format: {extension}")
        
        try:
            if extension == '.pdf':
                content = self._load_pdf_from_bytes(file_bytes)
            elif extension == '.txt':
                content = self._load_text_from_bytes(file_bytes)
            elif extension in ['.docx', '.doc']:
                content = self._load_docx_from_bytes(file_bytes)
            else:
                raise ValueError(f"Unsupported file format: {extension}")
            
            file_size = len(file_bytes.getvalue())
            
            return {
                'filename': filename,
                'content': content,
                'file_size': file_size,
                'file_size_formatted': format_file_size(file_size),
                'extension': extension,
                'char_count': len(content),
                'word_count': len(content.split()),
            }
            
        except Exception as e:
            self.logger.error(f"Error loading document {filename}: {str(e)}")
            raise
    
    def _load_pdf(self, file_path: Path) -> str:
        """Load text from PDF file."""
        text_content = []
        
        try:
            # Try with pypdf first (more modern)
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(text)
                    except Exception as e:
                        self.logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        
        except Exception:
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            text = page.extract_text()
                            if text.strip():
                                text_content.append(text)
                        except Exception as e:
                            self.logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
            except Exception as e:
                raise ValueError(f"Failed to read PDF file: {str(e)}")
        
        if not text_content:
            raise ValueError("No text content could be extracted from the PDF")
        
        return '\n\n'.join(text_content)
    
    def _load_pdf_from_bytes(self, file_bytes: io.BytesIO) -> str:
        """Load text from PDF BytesIO object."""
        text_content = []
        
        try:
            pdf_reader = pypdf.PdfReader(file_bytes)
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)
                except Exception as e:
                    self.logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                    
        except Exception:
            # Fallback to PyPDF2
            try:
                file_bytes.seek(0)  # Reset position
                pdf_reader = PyPDF2.PdfReader(file_bytes)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(text)
                    except Exception as e:
                        self.logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
            except Exception as e:
                raise ValueError(f"Failed to read PDF file: {str(e)}")
        
        if not text_content:
            raise ValueError("No text content could be extracted from the PDF")
        
        return '\n\n'.join(text_content)
    
    def _load_text(self, file_path: Path) -> str:
        """Load text from plain text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode text file with any supported encoding")
            
        except Exception as e:
            raise ValueError(f"Failed to read text file: {str(e)}")
    
    def _load_text_from_bytes(self, file_bytes: io.BytesIO) -> str:
        """Load text from BytesIO object."""
        try:
            content = file_bytes.getvalue()
            
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode text file with any supported encoding")
            
        except Exception as e:
            raise ValueError(f"Failed to read text file: {str(e)}")
    
    def _load_docx(self, file_path: Path) -> str:
        """Load text from DOCX file."""
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            if not text_content:
                raise ValueError("No text content found in the document")
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            raise ValueError(f"Failed to read DOCX file: {str(e)}")
    
    def _load_docx_from_bytes(self, file_bytes: io.BytesIO) -> str:
        """Load text from DOCX BytesIO object."""
        try:
            doc = DocxDocument(file_bytes)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            if not text_content:
                raise ValueError("No text content found in the document")
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            raise ValueError(f"Failed to read DOCX file: {str(e)}")
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return list(self.SUPPORTED_EXTENSIONS)
